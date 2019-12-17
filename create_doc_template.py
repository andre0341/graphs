# -*- coding: utf-8 -*-
# !/usr/bin/env
# !/Library/Frameworks/Python.framework/Versions/3.8/bin/python3
import xlrd
from docx import Document
from docx.shared import Inches

def createdoctemplate(d):
    document = Document()

    # start writing docx file
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "Report Form\tPMI-Network\t" + d['Nome']
    paragraph.style = document.styles["Header"]


    # ANAGRAFICA e SETTORE
    document.add_heading('PMI - Network ', 0)
    p = document.add_paragraph('L\'azienda ')
    p.add_run(str(d['Nome'])).bold = True
    p.add_run (' opera nel settore:' + str(d['Qual è il tuo settore di riferimento?']) + '')

    # DIGITALIZZAZIONE
    document.add_heading('Digitalizzazione dell\' azienda ', level=1)
    digit = document.add_paragraph('L\'azienda presenta un discreto grado di digitalizzazione. \n')
    digit.add_run('In relazione alle tecnologie 4.0 l\'azienda presenta..')
    document.add_picture('tecnologie40.png', width=Inches(6))

    # STRATEGIA
    document.add_heading('Strategia  Aziendale', level=1)
    strat = document.add_paragraph(str(d['Nome']))
    strat.add_run(' ha una chiara strategia per i prossimi anni.\n')
    strat.add_run('' + '' + str(d['Priorità/Necessità: Altro']) + '')
    document.add_picture('produzione.png', width=Inches(6))
    document.add_picture('ricerca_e_innovazione.png', width=Inches(6))
    document.add_picture('gestione_aziendale.png', width=Inches(6))

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')
    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    # document.add_picture('radar.png', width=Inches(1.25))
    
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    '''table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    '''
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True

    document.add_page_break()

    document.save(str(d['Nome']) + '_demo.docx')
