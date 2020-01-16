# -*- coding: utf-8 -*-
# !/usr/bin/env
# !/Library/Frameworks/Python.framework/Versions/3.8/bin/python3
import xlrd
from docx import Document
from docx.shared import Inches, RGBColor

def createdoctemplate(d):
    document = Document()

    # start writing docx file
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "Report Form\tPMI-Network\t" + d['Nome']
    paragraph.style = document.styles["Header"]


    # ANAGRAFICA e SETTORE
    document.add_heading('PMI - Network ', 0)
    p = document.add_paragraph('L\'azienda ')
    p.add_run(str(d['Nome'])).bold = True
    p.add_run (' opera nel settore: ' + str(d['Qual è il tuo settore di riferimento?']) + '')

    # DIGITALIZZAZIONE
    document.add_heading('Digitalizzazione dell\' azienda ', level=1)
    digit = document.add_paragraph('L\'azienda presenta un discreto grado di digitalizzazione. \n')
    digit.add_run('In relazione alle tecnologie 4.0 l\'azienda presenta..')
    document.add_picture('tecnologie40.png', width=Inches(6))

    # STRATEGIA
    document.add_heading('Strategia  Aziendale', level=1)
    strat = document.add_paragraph(str(d['Nome']))
    strat.add_run(' ha una chiara strategia per i prossimi anni.\n')
    strat.add_run('' + '' + str(d['Priorità/Necessità: Altro']) + '')
    document.add_picture('produzione.png', width=Inches(6))
    document.add_picture('ricerca_e_innovazione.png', width=Inches(6))
    document.add_picture('gestione_aziendale.png', width=Inches(6))
    
    document.add_page_break() 
    
    document.add_heading('Risposte al Questionario', level=1)
    document.add_paragraph('Di seguito l\'elenco completo delle risposte fornite in data')
    document.add_paragraph(' INSERIRE_DATA_INCONTRO :')

    #Domande
    #Risposta
        
    '''p = document.add_paragraph()
    title_of_question = p.add_run('I clienti dell\'impresa sono:')
    title_of_question.bold = True
    title_of_question.italic = True'''
    
    # Print All the values in dictionary
    '''for i in d : 
        print("%s  %s" %(i, d[i]))'''

    for i in d:
        #document.add_paragraph(i)
        p = document.add_paragraph()
        question = p.add_run(i)
        question.bold = True
        question.italic = True

        stringa = str(d[i])
        document.add_paragraph(stringa).bold = False

    '''
    document.add_paragraph('Il vantaggio competitivo dell\'azienda è:').bold = True
    document.add_paragraph('In azienda è presente un\'attività di formazione permanente del personale?')
    document.add_paragraph('Numero di laureati presenti in azienda')
    document.add_paragraph('Numero di consulenti esterni di tipo tecnico (numero di persone, no commercialista o avvocato)')
    '''

    document.add_page_break() 
    document.add_heading('Contatti e Next Step', level=1)
    document.add_paragraph('Per prenotare un nuovo incontro con un esperto di progetto o per eeventuali dubbi, invitiamo a contattare il gruppo di progetto PMI-Network tramite:') 
    document.add_paragraph('posta elettronica ') 
    document.add_paragraph().add_run('pmi-network@polimi.it').font.color.rgb = RGBColor(0x42, 0x24, 0xE9)



    # document.add_picture('radar.png', width=Inches(1.25))
    document.save(str(d['Nome']) + '_demo.docx')
